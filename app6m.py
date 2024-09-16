# Works perfect except the Light Dark and Blue background colours

import streamlit as st
import os
import openai
import pandas as pd
from docx import Document
from langchain_community.document_loaders import Docx2txtLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.chat_models import ChatOpenAI
from langchain.vectorstores import Qdrant
from langchain.chains import RetrievalQA
from langchain import hub
from dotenv import load_dotenv
import pyttsx3  # For text-to-speech functionality

# Load environment variables from .env file
load_dotenv()

# Set OpenAI API key
openai.api_key = os.getenv("MY_KEY")

# Load the latest version of the prompt
prompt = hub.pull("rlm/rag-prompt", api_url="https://api.hub.langchain.com")

# Load the .docx document
docx_loader = Docx2txtLoader(r"C:\Users\dhaka\Desktop\AI Biology Tutor\NCERT Biology\Combined_Chapters.docx")
raw_documents = docx_loader.load()
text_splitter = CharacterTextSplitter(separator=" ", chunk_size=1000, chunk_overlap=200)
documents = text_splitter.split_documents(raw_documents)

# Store splits in vector store
vectorstore = Qdrant.from_documents(documents=documents, embedding=OpenAIEmbeddings(openai_api_key=openai.api_key))

# Language Model
llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0, openai_api_key=openai.api_key)

# RetrievalQA
qa_chain = RetrievalQA.from_chain_type(
    llm, retriever=vectorstore.as_retriever(), chain_type_kwargs={"prompt": prompt}
)

# Function to get learning style from Excel
def get_learning_style(name):
    df = pd.read_excel('student_details.xlsx')
    student_row = df[df['Name'].str.lower() == name.lower()]
    if not student_row.empty:
        return student_row.iloc[0]['Learning Style'].strip().lower()
    return None

# Function to generate specific feedback for the question based on learning style
def generate_specific_feedback(question, correct_answer, learning_style, level):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant providing specific strategies for learning."},
            {"role": "user", "content": f"Given that the learner has a {learning_style} learning style, what specific strategy would you suggest for the following question that was answered incorrectly?\n\nQuestion: {question}\nCorrect Answer: {correct_answer}\nBloom's Taxonomy Level: {level}\n\nPlease provide a strategy that aligns with the {learning_style} learning style."}
        ]
    )
    strategy = response['choices'][0]['message']['content']
    return strategy

def generate_summary_based_on_style(content, style):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": f"You are a helpful assistant that summarizes content based on different learning styles. The style is {style}."},
            {"role": "user", "content": f"Please summarize the following content with an emphasis on {style} learning style:\n\n{content}"}
        ]
    )
    summary = response['choices'][0]['message']['content']
    return summary
def handle_chat():
    user_message = st.session_state.user_input.strip()
    if user_message:
        # Get response from the chat model
        response_data = qa_chain({"query": user_message})
        # Check if result is available and set the response
        if response_data.get("result") and response_data["result"].strip():
            response = response_data["result"]
        else:
            response = "I don't know."
        
        # Update chat history
        st.session_state.chat_history.append(("You", user_message))
        st.session_state.chat_history.append(("Bot", response))

        # Clear the input box after processing
        st.session_state.user_input = ""
# Fixing theme issue
def apply_theme(selected_theme):
    color_options = {
        'Light': {'background': '#FFFFFF', 'text': '#000000'},
        'Dark': {'background': '#000000', 'text': '#FFFFFF'},
        'Blue': {'background': '#DFF6FF', 'text': '#000000'},
    }
    color_scheme = color_options.get(selected_theme, color_options['Light'])
    st.markdown(
        f"""
        <style>
        .reportview-container {{
            background-color: {color_scheme['background']};
            color: {color_scheme['text']};
        }}
        .sidebar .sidebar-content {{
            background-color: {color_scheme['background']};
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

# Text-to-speech function
def text_to_speech(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()

# Path to the NCERT Biology folder
folder_path = r'C:\Users\dhaka\Desktop\AI Biology Tutor\NCERT Biology'

# Get the list of chapters (subfolder names) from the NCERT Biology folder
chapters = [name for name in os.listdir(folder_path) if os.path.isdir(os.path.join(folder_path, name))]

# List of options for the selected chapter
options = ["None", "Summary", "Exercise Help", "Quiz", "Chatbot"]

# List of languages GPT-3.5 Turbo can translate into, including Indian languages and English
languages = [
    "None", "English", "French", "Spanish", "German", "Chinese", "Japanese", "Hindi", 
    "Arabic", "Russian", "Portuguese", "Italian", "Bengali", "Tamil", 
    "Telugu", "Marathi", "Gujarati", "Kannada", "Malayalam", "Punjabi"
]

# Learning styles available
learning_styles = [
    "None", "Explanatory", "Socratic (Probing)", "Adaptive", "Active recall (Questionnaire)",
    "Analogy based", "Story telling"
]

st.title("Interactive Learning Platform")

# Sidebar for theme selection and accessibility
st.sidebar.write("Accessibility Options")
selected_theme = st.sidebar.selectbox("Select a Theme", ['Light', 'Dark', 'Blue'])
apply_theme(selected_theme)

# Font size selector
font_size = st.sidebar.slider("Adjust Font Size", 12, 30, 16)

# Text-to-speech option
enable_tts = st.sidebar.checkbox("Enable Text-to-Speech")

# Dropdown for selecting a chapter
st.write("Select a chapter from the dropdown")
selected_chapter = st.selectbox("Select a Chapter", chapters)

# Display the options dropdown only after a chapter is selected
if selected_chapter:
    selected_option = st.selectbox("Choose an Option", options)
    
    if selected_option == "Summary":
        selected_style = st.selectbox("Choose your learning style", learning_styles)
        
        if selected_style:
            # Load the summary content from the document
            summary_file_name = f"{selected_chapter} Summary.docx"
            summary_path = os.path.join(folder_path, selected_chapter, summary_file_name)
            
            if os.path.exists(summary_path):
                doc = Document(summary_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
                
                # Generate the summary based on the selected style
                styled_summary = generate_summary_based_on_style(full_text, selected_style)
                
                # Display the styled summary with adjustable font size
                st.markdown(f"<p style='font-size:{font_size}px;'>{styled_summary}</p>", unsafe_allow_html=True)
                
                # Option for text-to-speech
                if enable_tts:
                    text_to_speech(styled_summary)

                # Ask for translation
                selected_language = st.selectbox("Select a Language for Translation", languages)
                
                if selected_language and selected_language != "None":
                    # Translate the styled summary
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant that translates text."},
                            {"role": "user", "content": f"Translate the following text to {selected_language}:\n\n{styled_summary}"}
                        ]
                    )
                    
                    # Extract translated text from the response
                    translated_text = response['choices'][0]['message']['content']
                    
                    # Display the translated summary
                    st.markdown(f"<p style='font-size:{font_size}px;'>{translated_text}</p>", unsafe_allow_html=True)

                    if enable_tts:
                        text_to_speech(translated_text)
            else:
                st.error(f"{summary_file_name} file not found for {selected_chapter}.")
    
    elif selected_option == "Exercise Help":
        # Path to the Exercises.docx file with the expected format
        exercises_file_name = f"{selected_chapter} Exercises.docx"
        exercises_path = os.path.join(folder_path, selected_chapter, exercises_file_name)
        
        if os.path.exists(exercises_path):
            # Load the Exercises.docx content
            doc = Document(exercises_path)
            questions = ["None"] + [para.text for para in doc.paragraphs if para.text.strip() != ""]
            
            if questions:
                selected_question = st.selectbox("Select an Exercise Question", questions)
                
                if selected_question and selected_question != "None":
                    # Send the selected question to GPT-3.5 Turbo for a response
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant that provides concise answers to questions."},
                            {"role": "user", "content": f"Provide a 4-5 sentence answer to the following question:\n\n{selected_question}"}
                        ]
                    )
                    
                    # Extract the answer from the response
                    answer = response['choices'][0]['message']['content']
                    
                    # Display the answer on Streamlit
                    st.write(f"### Answer to the Question")
                    st.write(answer)
                    
                    # Ask if the user wants translation
                    translate = st.selectbox("Would you like to translate the question and answer?", ["No", "Yes"])
                    
                    if translate == "Yes":
                        translation_language = st.selectbox("Select a Language for Translation", languages)
                        
                        if translation_language and translation_language != "None":
                            # Translate the question and answer
                            translation_response = openai.ChatCompletion.create(
                                model="gpt-3.5-turbo",
                                messages=[
                                    {"role": "system", "content": "You are a helpful assistant that translates text."},
                                    {"role": "user", "content": f"Translate the following text to {translation_language}:\n\nQuestion: {selected_question}\n\nAnswer: {answer}"}
                                ]
                            )
                            
                            translated_text = translation_response['choices'][0]['message']['content']
                            
                            # Display the question and answer in both languages
                            st.write(f"### Question and Answer in {translation_language}")
                            st.write(translated_text)
                            
                            if enable_tts:
                                text_to_speech(translated_text)
            else:
                st.error(f"No questions found in {exercises_file_name}.")
        else:
            st.error(f"{exercises_file_name} file not found for {selected_chapter}.")



    elif selected_option == "Chatbot":
        # Initialize chat history if not present
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []

        st.title('Welcome to the Biology Chatbot')
        st.write("Ask your questions here!")

        # Display chat history
        for role, message in st.session_state.chat_history:
            st.text_area(f"{role}:", value=message, height=100, max_chars=None)

        # UI for sending new messages
        user_input = st.text_input("Type your question about Biology here:", key="user_input")
        send_button = st.button("Send", on_click=handle_chat)

    
    
    elif selected_option == "Quiz":
        # Path to the Summary and Exercises files
        summary_file_name = f"{selected_chapter} Summary.docx"
        exercises_file_name = f"{selected_chapter} Exercises.docx"
        summary_path = os.path.join(folder_path, selected_chapter, summary_file_name)
        exercises_path = os.path.join(folder_path, selected_chapter, exercises_file_name)

        if os.path.exists(summary_path) and os.path.exists(exercises_path):
            # Load the content from both the Summary and Exercises files
            summary_doc = Document(summary_path)
            exercises_doc = Document(exercises_path)
            
            summary_text = "\n".join([para.text for para in summary_doc.paragraphs])
            exercises_text = "\n".join([para.text for para in exercises_doc.paragraphs if para.text.strip() != ""])
            
            # Bloom's Taxonomy levels
            blooms_levels = [
                "Remember",    # Lowest level
                "Understand",
                "Apply",
                "Analyze",
                "Evaluate"     # Highest level for this quiz
            ]
            
            # Initialize session state variables
            if 'current_question_idx' not in st.session_state:
                st.session_state.current_question_idx = 0
                st.session_state.score = 0
                st.session_state.quiz_complete = False
                st.session_state.correct_answer = ""
                st.session_state.user_choice = None
                st.session_state.question = ""
                st.session_state.choices = []
                st.session_state.incorrect_questions = []  # To track incorrect questions
                st.session_state.levels_mastered = [False] * len(blooms_levels)  # Track mastery for each level

            current_question_idx = st.session_state.current_question_idx
            
            if not st.session_state.quiz_complete and current_question_idx < len(blooms_levels):
                current_level = blooms_levels[current_question_idx]
                
                if st.session_state.user_choice is None:  # Generate question only if not already answered
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": f"You are a helpful assistant that generates questions based on Bloom's Taxonomy at the {current_level} level."},
                            {"role": "user", "content": f"Generate a multiple-choice question with four options from the following content:\n\nSummary: {summary_text}\n\nExercises: {exercises_text}\n\nLevel: {current_level}. Please provide the correct answer at the end in the format 'Correct Answer: [answer]'."}
                        ]
                    )
                    
                    quiz_content = response['choices'][0]['message']['content'].split('\n')
                    quiz_content = [line for line in quiz_content if line.strip()]  # Filter out empty lines
                    
                    if len(quiz_content) >= 3:  # Ensure there are at least a question and two choices
                        st.session_state.question = quiz_content[0]
                        st.session_state.choices = quiz_content[1:-1]
                        correct_answer_line = quiz_content[-1]
                        
                        if 'Correct Answer:' in correct_answer_line:
                            try:
                                st.session_state.correct_answer = correct_answer_line.split(': ', 1)[1]
                            except IndexError:
                                st.error("Could not determine the correct answer from the generated content.")
                                st.write("Full Response:")
                                st.write("\n".join(quiz_content))
                        else:
                            st.error("Could not determine the correct answer from the generated content.")
                            st.write("Full Response:")
                            st.write("\n".join(quiz_content))
                    else:
                        st.error("Generated quiz content is incomplete.")
                        st.write("Full Response:")
                        st.write("\n".join(quiz_content))
                
                # Display the question and options
                st.write(f"### Question {current_question_idx + 1} ({current_level} Level)")
                st.write(st.session_state.question)
                
                st.session_state.user_choice = st.radio("Select your answer:", st.session_state.choices, index=0)
                
                # Submit button to check answer
                if st.button("Submit"):
                    if st.session_state.user_choice == st.session_state.correct_answer:
                        st.write("Good job! That's correct.")
                        st.session_state.score += 1
                        st.session_state.levels_mastered[current_question_idx] = True  # Mark level as mastered
                    else:
                        st.write(f"Incorrect. The correct answer is: {st.session_state.correct_answer}")
                        # Track incorrect question details
                        st.session_state.incorrect_questions.append({
                            "question": st.session_state.question,
                            "correct_answer": st.session_state.correct_answer,
                            "level": current_level
                        })
                    
                    # Prepare for the next question
                    st.session_state.current_question_idx += 1
                    st.session_state.user_choice = None  # Reset for next question
                    st.session_state.question = ""  # Clear the current question
                    st.session_state.choices = []  # Clear the choices
                    
                    if st.session_state.current_question_idx >= len(blooms_levels):
                        st.session_state.quiz_complete = True
                    
                    st.experimental_rerun()  # Rerun to load the next question
                
            if st.session_state.quiz_complete:
                st.write(f"Quiz completed! You scored {st.session_state.score} out of {len(blooms_levels)}")
    
                # Display a report on mastered and incorrect levels
                student_name = "Padma"  # Replace with dynamic input as needed
                learning_style = get_learning_style(student_name)
                
                for idx, level in enumerate(blooms_levels):
                    if st.session_state.levels_mastered[idx]:
                        st.write(f"{level} level - mastered!")
                    else:
                        st.write(f"You need to review the content again for the {level} level.")
                        
                        # Provide specific feedback based on the incorrect question
                        for question_detail in st.session_state.incorrect_questions:
                            if question_detail["level"] == level:
                                feedback = generate_specific_feedback(question_detail["question"], question_detail["correct_answer"], learning_style, level)
                                st.write(f"Suggested Strategy ({learning_style.capitalize()} Learner): {feedback}")
                    
                st.write("Please revise the incorrect levels if any and try the quiz again.")
    
                # Button to return to the original choice screen
                if st.button("Retake the Quiz"):
                    # Clear only the quiz-related session state variables
                    st.session_state.current_question_idx = 0
                    st.session_state.score = 0
                    st.session_state.quiz_complete = False
                    st.session_state.correct_answer = ""
                    st.session_state.user_choice = None
                    st.session_state.question = ""
                    st.session_state.choices = []
                    st.session_state.incorrect_questions = []  # Reset incorrect questions
                    st.session_state.levels_mastered = [False] * len(blooms_levels)  # Reset level mastery
                    st.experimental_rerun()
        else:
            st.error(f"Summary or Exercises files not found for {selected_chapter}.")
